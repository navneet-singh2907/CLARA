"""Build the CLARA Week 4 submission documentation as a Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "CLARA_Week4_Submission_Documentation.docx"
FALLBACK_OUT_PATH = ROOT / "CLARA_Week4_Submission_Documentation_updated.docx"


BRAND = "9F463E"
INK = "111827"
MUTED = "667085"
LIGHT = "F4F7FB"
GREEN = "DDEFE5"
AMBER = "FFF1D6"
RED = "FCE4E2"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(INK)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()


def add_h1(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_h2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BRAND)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT) -> None:
    paragraph = doc.add_paragraph()
    shade_paragraph(paragraph, fill)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    run.font.size = Pt(10.5)
    paragraph.add_run("\n")
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor.from_string(INK)
    run.font.size = Pt(10)


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    shade_paragraph(paragraph, "F3F4F6")
    for line in code.strip("\n").splitlines():
        run = paragraph.add_run(line.rstrip() + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("1F2937")
    paragraph.paragraph_format.space_after = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(header_cells[idx], BRAND)
        if widths:
            header_cells[idx].width = Inches(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def build_doc() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    add_title(
        doc,
        "CLARA - Week 4 AI Evals Submission Documentation",
        "Credit Loan Analysis & Review Agent | LangChain + LangGraph + FastAPI SSE + Next.js + LangSmith + MCP",
    )

    add_callout(
        doc,
        "Executive summary",
        "CLARA started as a Week 3 multi-agent small business loan review pipeline. "
        "For Week 4, it was hardened into an evaluation-first agentic AI system: "
        "50-case gold set, targeted regression loops, LangSmith traces, MCP evidence tools, "
        "agent-level failure attribution, live drift probes, judge agreement, PDF packets, "
        "and production guardrails for LLM failures and uploads.",
        GREEN,
    )

    add_table(
        doc,
        ["Item", "Submission Detail"],
        [
            ["Product", "CLARA - Credit Loan Analysis & Review Agent"],
            ["Domain", "Small business loan application review"],
            ["Core agents", "Term Extractor, Compliance Checker, Credit Risk Scorer, Review Synthesizer"],
            ["Evaluation", "50-case gold set, failure attribution, before/after delta report, judge agreement, drift probe"],
            ["Observability", "FastAPI SSE live timeline, LangSmith traces, MCP inspection server"],
            ["Current status", "Dockerized, CI-tested, local and Vercel-ready architecture"],
        ],
        [1.5, 5.6],
    )

    add_h1(doc, "Evaluation Framing")
    add_callout(
        doc,
        "Worked example",
        "I measure final loan-decision accuracy, compliance red-flag detection, "
        "risk-band calibration, agent trajectory correctness, judge agreement, drift stability, "
        "and p95 latency on CLARA, a multi-agent small business loan review pipeline, using a "
        "frozen 50-case golden dataset covering clean, ambiguous, edge, known-failure, and "
        "adversarial loan scenarios. I use exact-match and code-based evaluators for structured "
        "labels, LLM-as-judge for faithfulness and explainability, and primary-vs-secondary judge "
        "agreement for evaluator reliability. Pass bar: >=90% final outcome accuracy, >=90% "
        "compliance accuracy, >=85% risk-band accuracy, 0 critical false approvals, 100% trace "
        "coverage for live runs, and p95 latency low enough for demo review. I run this in "
        "LangSmith and report the measured delta from baseline to post-improvement.",
        GREEN,
    )
    add_table(
        doc,
        ["Evaluation Decision", "CLARA Answer"],
        [
            ["Agent under test", "CLARA multi-agent small business loan review pipeline."],
            ["User outcome", "A loan officer needs a correct approve, reject, or escalate recommendation with evidence they can audit."],
            ["Metrics", "Outcome accuracy, compliance accuracy, risk-band accuracy, trajectory correctness, judge agreement, drift stability, p95 latency."],
            ["Judge method", "Exact-match and code-based evaluators for structured labels; LLM-as-judge and primary/secondary judge agreement for qualitative review."],
            ["Golden dataset", "Frozen 50-case SBA-style set with dataset version clara-week4-v1."],
            ["Pass bar", ">=90% final outcome, >=90% compliance, >=85% risk-band, 0 critical false approvals, 100% trace coverage."],
            ["Instrumentation", "LangSmith traces, SSE timeline, MCP evidence server, local artifacts, before/after dashboard."],
        ],
        [1.8, 5.3],
    )

    add_h1(doc, "1. From Week 3 Pipeline to Week 4 Eval Lab")
    add_body(
        doc,
        "The first version of CLARA solved the workflow problem: take a loan case, run it through specialist agents, "
        "and produce an approval, rejection, or escalation packet. Week 4 changed the question. The goal became: "
        "can we prove, inspect, and improve the system one failure at a time?",
    )
    add_bullet(doc, "Week 3: build the multi-agent review pipeline and Streamlit/Next.js demo.")
    add_bullet(doc, "Week 4: add rigorous AI evaluation, failure attribution, LangSmith tracing, MCP access, and targeted fixes.")
    add_bullet(doc, "Final result: a project that is not only functional, but inspectable and improvable.")

    add_h1(doc, "2. Week 4 Requirement Coverage")
    add_table(
        doc,
        ["Requirement", "How CLARA Covers It", "Evidence"],
        [
            ["Gold set", "Expanded from 30 to 50 cases with clean, ambiguous, adversarial, edge, and known-failure cases.", "loan_pipeline/data/gold_set.json"],
            ["Metrics", "Exact match, tier accuracy, failure category, agent attribution, judge agreement, drift fingerprints.", "loan_pipeline/eval/*"],
            ["Improvement loop", "Baseline vs improved regression table shows targeted fix impact.", "week4_compare.py"],
            ["LLM-as-judge", "Primary and secondary judge scoring with agreement and dimension deltas.", "inter_rater.py"],
            ["Observability", "LangSmith traces plus SSE timeline for live agent execution.", "api/streaming.py"],
            ["MCP", "MCP evidence server exposes gold-case comparison and trace inspection tools.", "mcp/evidence_server.py"],
            ["Documentation", "This document includes architecture, failures, fixes, snippets, and verification path.", "Submission artifact"],
        ],
        [1.35, 3.85, 1.9],
    )

    add_h1(doc, "3. Architecture")
    add_body(
        doc,
        "CLARA uses LangGraph as the orchestration layer. The term extractor runs first, the schema validator checks "
        "structured terms, then compliance and risk agents run as specialist reviewers. The synthesizer combines outputs, "
        "detects contradictions, creates counterfactual explanations, and prepares a human-reviewable packet.",
    )
    add_code(
        doc,
        """
Loan Case / Uploaded PDF
    |
    v
Term Extractor Agent
    |
    v
Schema Validator
    |
    +--> Compliance Checker Agent
    |
    +--> Credit Risk Scorer Agent
            |
            v
Review Synthesizer
    |       |       |       |
    |       |       |       +--> Counterfactual explanations
    |       |       +----------> Contradiction detection
    |       +------------------> Human override audit log
    +--------------------------> PDF packet + evaluation harness
        """,
    )

    add_h1(doc, "4. Agent Responsibilities")
    add_table(
        doc,
        ["Agent", "Graph Node", "Responsibility", "Failure Signal"],
        [
            ["Term Extractor", "term_extractor", "Extracts borrower, amount, SBA guarantee, term, credit score, missing docs, prior default, and rationale.", "Low term confidence, missing required fields, invalid JSON from LLM."],
            ["Schema Validator", "schema_validator", "Normalizes and validates the extracted loan case before specialist review.", "Validation error in execution trace."],
            ["Compliance Checker", "compliance_checker", "Applies SBA-style, bank, or CDFI policy rules and flags documentation or eligibility issues.", "Fallback compliance result, FAIL status, agent_errors entry."],
            ["Credit Risk Scorer", "credit_risk_scorer", "Scores risk band, risk score, confidence, and rationale.", "Risk mismatch, score drift, fallback risk result."],
            ["Review Synthesizer", "review_synthesizer", "Combines all agent outputs, detects contradictions, creates counterfactuals, and recommends outcome.", "Contradictions, missing packet fields, downstream escalation."],
            ["Judge Agents", "primary/secondary judge", "Evaluate packet faithfulness, completeness, risk calibration, compliance accuracy, and explainability.", "Low exact agreement, large dimension deltas."],
        ],
        [1.25, 1.25, 3.25, 1.75],
    )

    add_h1(doc, "5. What Went Wrong and How We Fixed It")

    add_h2(doc, "Fix 1 - Context-free LLM exceptions")
    add_body(
        doc,
        "Problem: when an LLM returned malformed JSON, the system originally surfaced a generic error like "
        "'LLM response must be valid JSON.' That was not enough to debug a multi-agent pipeline because it did not say "
        "which agent, which case, which provider, or which model failed.",
    )
    add_body(doc, "Fix: introduce a structured LLMResponseError carrying agent, case, operation, provider, model, temperature, and a safe response preview.")
    add_code(
        doc,
        """
class LLMResponseError(Exception):
    def __init__(
        self,
        message: str,
        *,
        agent_name: str,
        case_id: str | None,
        operation: str,
        provider: str,
        model: str,
        temperature: float,
        response_preview: str | None = None,
    ) -> None:
        self.agent_name = agent_name
        self.case_id = case_id
        self.operation = operation
        self.provider = provider
        self.model = model
        self.temperature = temperature
        self.response_preview = response_preview

        super().__init__(
            f"[{agent_name}|{operation}|case={case_id}|{provider}/{model}] {message}"
            + (f" | response_preview={response_preview!r}" if response_preview else "")
        )
        """,
    )
    add_callout(
        doc,
        "Impact",
        "A failure can now be traced to a specific agent and case instead of appearing as a generic pipeline crash.",
        GREEN,
    )

    add_h2(doc, "Fix 2 - LLM calls could hang indefinitely")
    add_body(
        doc,
        "Problem: live LLM calls had no timeout. If Nebius or OpenAI slowed down, the SSE stream could look frozen and the demo would appear broken.",
    )
    add_body(doc, "Fix: add a 30-second timeout directly to the ChatOpenAI client and keep API keys typed as SecretStr.")
    add_code(
        doc,
        """
from pydantic import SecretStr
from langchain_openai import ChatOpenAI

LLM_TIMEOUT_SECONDS = 30.0

llm = ChatOpenAI(
    api_key=SecretStr(api_key),
    base_url=settings.llm_base_url,
    model=model,
    temperature=temperature,
    timeout=LLM_TIMEOUT_SECONDS,
)
        """,
    )
    add_callout(
        doc,
        "Impact",
        "Live mode now fails fast enough to show a clear error event instead of leaving the UI stuck.",
        GREEN,
    )

    add_h2(doc, "Fix 3 - Invalid JSON was hard to inspect")
    add_body(
        doc,
        "Problem: JSON parsing failures did not expose enough of the model response to understand whether the problem was markdown, prose, missing braces, or an unrelated answer.",
    )
    add_body(doc, "Fix: capture a bounded response preview and attach it to the structured error.")
    add_code(
        doc,
        """
preview = content[:LLM_RESPONSE_PREVIEW_CHARS].replace("\\n", " ")

try:
    return json.loads(content)
except json.JSONDecodeError as exc:
    raise LLMResponseError(
        "Response is not valid JSON.",
        agent_name=agent_name,
        case_id=case_id,
        operation=operation,
        provider=provider,
        model=model,
        temperature=temperature,
        response_preview=preview,
    ) from exc
        """,
    )

    add_h2(doc, "Fix 4 - One agent failure could collapse the pipeline")
    add_body(
        doc,
        "Problem: if Compliance Checker or Credit Risk Scorer failed, the pipeline could terminate even though another agent still had usable evidence.",
    )
    add_body(
        doc,
        "Fix: isolate specialist-agent failures. The failed agent now returns a fallback result, records an ERROR trace entry, "
        "and lets the synthesizer continue with a human escalation signal.",
    )
    add_code(
        doc,
        """
except Exception as exc:
    error_message = f"Compliance checker failed: {exc}"
    return {
        "compliance": _fallback_compliance_result(error_message),
        "agent_errors": [error_message],
        "execution_trace": [
            _trace_entry(
                node="compliance_checker",
                stage="compliance",
                started_at=started_at,
                status="ERROR",
                notes=error_message,
            )
        ],
    }
        """,
    )
    add_callout(
        doc,
        "Impact",
        "The system behaves more like a real review workflow: one specialist can fail, but the case can still be escalated with evidence.",
        GREEN,
    )

    add_h2(doc, "Fix 5 - Uploaded documents needed guardrails")
    add_body(
        doc,
        "Problem: uploaded PDF/text files were accepted without a clear size cap. A very large file could consume memory or make the serverless function unstable.",
    )
    add_body(doc, "Fix: add a 10 MB upload limit before document parsing.")
    add_code(
        doc,
        """
MAX_UPLOAD_BYTES = 10_000_000

async def _extract_upload_text(file: UploadFile) -> str:
    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"Uploaded file is too large. Maximum size is {MAX_UPLOAD_BYTES} bytes.",
        )
        """,
    )

    add_h2(doc, "Fix 6 - API errors needed structured payloads")
    add_body(
        doc,
        "Problem: API errors were useful to developers only if they could see backend logs. In a deployed demo, the frontend needed structured error details.",
    )
    add_body(doc, "Fix: wrap document review failures with endpoint, file, and case context.")
    add_code(
        doc,
        """
try:
    packet = run_pipeline(loan_case, review_policy=policy)
except Exception as exc:
    raise HTTPException(
        status_code=502,
        detail=error_payload(
            exc,
            endpoint="/review/document",
            file_name=file.filename,
            case_id=loan_case.case_id,
        ),
    ) from exc
        """,
    )

    add_h2(doc, "Fix 7 - MCP evidence server for external inspection")
    add_body(
        doc,
        "Problem: LangSmith shows traces, but Week 4 needed evaluator-grade tooling. We added MCP so an external client can inspect cases, compare outputs to gold labels, and pull graph traces.",
    )
    add_code(
        doc,
        """
{
    "name": "inspect_pipeline_trace",
    "description": "Run CLARA deterministically and return agent trace entries, errors, and final packet.",
    "inputSchema": {
        "type": "object",
        "properties": {
            "case_id": {"type": "string"},
            "dataset": {"type": "string"},
            "policy": {"type": "string"}
        },
        "required": ["case_id"]
    },
}
        """,
    )
    add_body(doc, "The MCP command is exposed through a project script:")
    add_code(
        doc,
        """
[project.scripts]
clara-mcp-evidence = "loan_pipeline.mcp.evidence_server:main"
        """,
    )

    add_h2(doc, "Fix 8 - Improvement needed to be visible as a before/after delta")
    add_body(
        doc,
        "Problem: saying 'we improved the system' is weak. The evaluator needs to see exactly which cases changed and whether the fix helped or hurt.",
    )
    add_body(doc, "Fix: add a Week 4 comparison report that runs baseline vs improved behavior and renders deltas.")
    add_code(
        doc,
        """
def render_week4_improvement_report(
    baseline_results: list[dict[str, Any]],
    improved_results: list[dict[str, Any]],
) -> str:
    baseline_summary = summarize_results(baseline_results)
    improved_summary = summarize_results(improved_results)
    deltas = compare_results(baseline_results, improved_results)
    ...
        """,
    )

    add_h2(doc, "Fix 9 - Failures needed agent attribution")
    add_body(
        doc,
        "Problem: a wrong outcome is not enough. In a multi-agent system, the key question is which agent caused the wrong outcome and how it propagated.",
    )
    add_body(doc, "Fix: failure attribution maps mismatches back to responsible agents and graph nodes.")
    add_code(
        doc,
        """
if not exact["risk_correct"]:
    attributions.append(
        _behavioral(
            responsible_agent="Credit Risk Scorer Agent",
            graph_node="credit_risk_scorer",
            expected=f"Risk band {gold.get('expected_risk_band', 'unknown')}",
            actual=f"Risk band {packet.get('risk_band', 'unknown')}",
            downstream_impact="May change approval, rejection, or escalation recommendation.",
        )
    )
        """,
    )

    add_h1(doc, "6. How to Inspect Which Agent Failed")
    add_body(doc, "There are four inspection layers. In a demo or debugging session, use them in this order:")
    add_table(
        doc,
        ["Layer", "Where to Look", "What It Tells You"],
        [
            ["Frontend timeline", "Loan Review tab", "Which graph node started, completed, or errored during a live run."],
            ["Activity logs", "Evaluation, Ablation, Drift, Judge Agreement tabs", "Which batch item is running and whether progress is advancing."],
            ["LangSmith", "CLARA project traces", "Raw prompt, model call, latency, errors, and nested runs."],
            ["MCP", "clara-mcp-evidence inspect_pipeline_trace", "Machine-readable trace, errors, and packet for a specific case."],
            ["Failure attribution", "Week 4 eval report", "Which agent is responsible for a mismatch against gold labels."],
        ],
        [1.45, 2.05, 3.65],
    )
    add_body(doc, "A typical debugging flow:")
    add_code(
        doc,
        """
1. Identify failing case from evaluation table, e.g. ADV-007.
2. Open the LangSmith trace for that run.
3. Check execution_trace for node status:
   - term_extractor
   - schema_validator
   - compliance_checker
   - credit_risk_scorer
   - review_synthesizer
4. If a node has status ERROR, inspect agent_errors.
5. If output is valid but wrong, inspect failure attribution.
6. Apply one targeted fix and rerun the before/after comparison.
        """,
    )

    add_h1(doc, "7. Evaluation Design")
    add_body(
        doc,
        "The evaluation set is intentionally structured instead of being a flat list of examples. That makes the results more diagnostic.",
    )
    add_table(
        doc,
        ["Case Type", "Purpose"],
        [
            ["Clean", "Straightforward applications the system should handle reliably."],
            ["Ambiguous", "Missing fields, conflicting terms, or unclear documents."],
            ["Adversarial", "Prior defaults, missing KYC, prompt-injection-like content, or misleading terms."],
            ["Edge", "Partial data, unusual loan terms, and borderline policy decisions."],
            ["Known failures", "Cases CLARA previously struggled with, used to prove targeted improvement."],
        ],
        [1.5, 5.6],
    )
    add_body(doc, "CLARA evaluates more than final outcome. It also measures:")
    add_bullet(doc, "Risk-band accuracy")
    add_bullet(doc, "Compliance correctness")
    add_bullet(doc, "Outcome correctness")
    add_bullet(doc, "Failure category distribution")
    add_bullet(doc, "Primary vs secondary judge agreement")
    add_bullet(doc, "Drift across repeated live LLM runs")
    add_bullet(doc, "Confidence calibration and disagreement patterns")

    add_h1(doc, "8. Current Project Status")
    add_callout(
        doc,
        "Current state",
        "CLARA is now an evaluation-first, multi-agent AI review system. It has live LLM mode, deterministic fallback mode, "
        "LangGraph orchestration, LangSmith tracing, MCP inspection, CI, Docker support, FastAPI SSE streaming, Next.js UI, "
        "PDF packets, upload review, judge agreement, drift probing, and targeted Week 4 improvement reporting.",
        GREEN,
    )
    add_table(
        doc,
        ["Capability", "Status"],
        [
            ["Multi-agent pipeline", "Implemented"],
            ["Parallel specialist review", "Implemented for compliance and risk branches"],
            ["Live LLM operation", "Implemented through OpenAI-compatible provider config"],
            ["Deterministic baseline", "Implemented for reproducible evaluation"],
            ["LangSmith tracing", "Implemented"],
            ["LangSmith before/after dashboard", "Implemented through scripts/week4_langsmith_dashboard.py"],
            ["MCP evidence server", "Implemented"],
            ["SSE live progress", "Implemented"],
            ["PDF review packets", "Implemented"],
            ["Dockerized stack", "Implemented"],
            ["CI workflow", "Implemented"],
            ["Production persistence", "Partial - audit log is demo/session oriented"],
        ],
        [2.4, 4.6],
    )

    add_h1(doc, "9. Remaining Limitations and Honest Next Steps")
    add_body(
        doc,
        "The project is intentionally transparent about remaining production gaps. These do not block the bootcamp submission, "
        "but they show where the system would go next in a real financial-services environment.",
    )
    add_bullet(doc, "Replace in-memory audit history with durable database persistence.")
    add_bullet(doc, "Move rate limiting to a shared store such as Redis for multi-instance deployments.")
    add_bullet(doc, "Expand from SBA-style seeded cases to a fully processed public SBA corpus.")
    add_bullet(doc, "Add lender-specific configurable risk rules outside code.")
    add_bullet(doc, "Add more red-team document uploads and prompt-injection test packets.")
    add_bullet(doc, "Use the LangSmith dashboard command during the demo to show baseline and improved runs side by side.")

    add_h1(doc, "10. Final Submission Positioning")
    add_body(
        doc,
        "CLARA is not just a loan-review demo. It is a case study in how to build, observe, evaluate, and improve an agentic AI system. "
        "The strongest part of the project is not only the final loan decision, but the fact that every decision can be inspected: "
        "which agent acted, what evidence it used, where it failed, how the failure was attributed, and how the fix changed the evaluation table.",
    )
    add_callout(
        doc,
        "Builder-of-the-week argument",
        "CLARA combines product realism with evaluation rigor: multi-agent workflow, live LLM variance checks, human override governance, "
        "judge agreement, MCP inspection, LangSmith traces, and targeted failure repair. It demonstrates the exact engineering behavior "
        "Week 4 is asking for: not just building agents, but proving and improving them.",
        AMBER,
    )

    add_h1(doc, "Appendix A - Useful Commands")
    add_code(
        doc,
        """
# Run full Python test suite
python -m pytest

# Run lint checks
ruff check loan_pipeline tests

# Run the API locally
uvicorn loan_pipeline.api.app:app --reload --port 8000

# Run the Next.js UI locally
cd web
npm run dev

# Run Dockerized stack
docker compose up --build

# Start MCP evidence server
clara-mcp-evidence

# Create LangSmith baseline-vs-improved dashboard
python scripts/week4_langsmith_dashboard.py
        """,
    )

    add_h1(doc, "Appendix B - Submission Checklist")
    add_bullet(doc, "Show the live UI with CLARA branding and system readiness.")
    add_bullet(doc, "Run one adversarial loan case and show the live agent timeline.")
    add_bullet(doc, "Download the PDF review packet.")
    add_bullet(doc, "Upload the packet to Judge Agreement and compare primary/secondary scores.")
    add_bullet(doc, "Run a live drift probe on a selected case.")
    add_bullet(doc, "Show LangSmith traces for a live LLM run.")
    add_bullet(doc, "Show MCP evidence inspection for a failing or improved case.")
    add_bullet(doc, "Show before/after Week 4 improvement report.")
    add_bullet(doc, "Explain the exact code fixes included in this document.")

    try:
        doc.save(OUT_PATH)
        return OUT_PATH
    except PermissionError:
        doc.save(FALLBACK_OUT_PATH)
        return FALLBACK_OUT_PATH


if __name__ == "__main__":
    print(build_doc())
